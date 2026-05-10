"""Builds shared correction plans and renders inline and hybrid DOCX outputs."""

import difflib
import copy
import re
from lxml import etree
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from toolkit.llm_service import get_corrections_from_llm, get_text_from_llm
from toolkit.prompts import get_prompt_max_input_words, DEFAULT_PROMPT_KEY


def _normalize_hidden_whitespace(text: str) -> tuple[str, int]:
    """Normalize invisible Unicode whitespace to regular spaces; return (normalized_text, count_of_replacements).
    
    Maps NBSP, narrow NBSP, figure space, thin space, and other invisible whitespace to regular space.
    Also removes zero-width characters (ZWSP, ZWNJ, ZWJ, BOM).
    """
    if not text:
        return text, 0
    
    original_text = text
    replacement_count = 0
    
    # Invisible space characters to replace with regular space
    invisible_spaces = {
        '\xa0': ' ',      # Non-breaking space (NBSP)
        '\u202f': ' ',    # Narrow no-break space
        '\u2000': ' ',    # En quad
        '\u2001': ' ',    # Em quad
        '\u2002': ' ',    # En space
        '\u2003': ' ',    # Em space
        '\u2004': ' ',    # Three-per-em space
        '\u2005': ' ',    # Four-per-em space
        '\u2006': ' ',    # Six-per-em space
        '\u2007': ' ',    # Figure space
        '\u2008': ' ',    # Punctuation space
        '\u2009': ' ',    # Thin space
        '\u200a': ' ',    # Hair space
    }
    
    for char, replacement in invisible_spaces.items():
        if char in text:
            text = text.replace(char, replacement)
            replacement_count += text.count(replacement) if replacement == ' ' else 1
    
    # Zero-width characters to remove
    zero_width_chars = {
        '\u200b',  # Zero-width space
        '\u200c',  # Zero-width non-joiner
        '\u200d',  # Zero-width joiner
        '\ufeff',  # Zero-width no-break space (BOM)
    }
    
    for char in zero_width_chars:
        if char in text:
            replacement_count += text.count(char)
            text = text.replace(char, '')
    
    return text, replacement_count


def _paragraph_contains_image(paragraph):
    """Returns True if the paragraph contains an inline/floating image."""
    p = paragraph._p
    return (
        p.find(f".//{qn('w:drawing')}") is not None
        or p.find(f".//{qn('w:pict')}") is not None
    )


def _insert_blank_line_before_images(doc):
    """Ensures a blank paragraph exists before and after every image paragraph."""
    inserted_count = 0
    for para in list(doc.paragraphs):
        if not _paragraph_contains_image(para):
            continue

        prev = para._p.getprevious()
        prev_is_blank_para = (
            prev is not None
            and prev.tag.endswith("}p")
            and "".join(prev.itertext()).strip() == ""
        )

        if not prev_is_blank_para:
            para.insert_paragraph_before("")
            inserted_count += 1

        nxt = para._p.getnext()
        next_is_blank_para = (
            nxt is not None
            and nxt.tag.endswith("}p")
            and "".join(nxt.itertext()).strip() == ""
        )

        if not next_is_blank_para:
            para._p.addnext(copy.deepcopy(para._p))
            new_next = para._p.getnext()
            for child in list(new_next):
                new_next.remove(child)
            inserted_count += 1

    return inserted_count


def _preserve_soft_breaks(original_text, corrected_text):
    """Reinsert original soft line breaks (Shift+Enter) if the correction dropped them."""
    break_char = "\n"
    original_breaks = original_text.count(break_char)
    if original_breaks == 0:
        return corrected_text
    if corrected_text.count(break_char) >= original_breaks:
        return corrected_text

    parts = original_text.split(break_char)
    if len(parts) <= 1:
        return corrected_text

    total_original_len = sum(len(p) for p in parts)
    if total_original_len <= 0:
        return corrected_text

    corrected_len = len(corrected_text)
    target_lengths = [round((len(p) / total_original_len) * corrected_len) for p in parts]
    diff = corrected_len - sum(target_lengths)
    if target_lengths:
        target_lengths[-1] += diff

    chunks = []
    cursor = 0
    for i, target in enumerate(target_lengths):
        remaining = corrected_len - cursor
        if i == len(target_lengths) - 1:
            chunk_len = remaining
        else:
            chunk_len = max(0, min(int(target), remaining))
        chunks.append(corrected_text[cursor:cursor + chunk_len])
        cursor += chunk_len

    return break_char.join(chunks)


def _build_deletion_marker(deleted_text, show_marker=True):
    """Builds a visible marker for deleted text so removals are obvious in output."""
    if not deleted_text or not show_marker:
        return ""
    visible = deleted_text.replace("\n", r"\n")
    return f"[-{visible}-]"


def _collect_all_paragraphs(doc):
    """Collect paragraphs from the document body and table cells."""
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    return all_paragraphs


def _find_all_indices(text, needle):
    """Return start indices of all non-overlapping needle matches in text."""
    indices = []
    if not needle:
        return indices
    start = 0
    while True:
        idx = text.find(needle, start)
        if idx == -1:
            break
        indices.append(idx)
        start = idx + max(1, len(needle))
    return indices


def _build_atomic_edits(original, corrected):
    """Split one sentence-level correction into atomic span edits."""
    edits = []
    matcher = difflib.SequenceMatcher(None, original, corrected)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue

        rel_start = i1
        orig_span = original[i1:i2]
        corr_span = corrected[j1:j2]

        if tag == 'insert':
            # Anchor inserts to a neighboring character so they can be located reliably.
            if i1 > 0:
                rel_start = i1 - 1
                anchor = original[rel_start:i1]
                orig_span = anchor
                corr_span = anchor + corr_span
            elif original:
                rel_start = 0
                anchor = original[0:1]
                orig_span = anchor
                corr_span = corr_span + anchor
            else:
                continue

        if not orig_span or orig_span == corr_span:
            continue

        edits.append(
            {
                'relative_start': rel_start,
                'original': orig_span,
                'corrected': corr_span,
            }
        )

    return edits


def _resolve_correction_start(block_content, original, preferred_start, search_start):
    """Resolve correction start index using an optional preferred offset hint."""
    if isinstance(preferred_start, int) and preferred_start >= search_start:
        end = preferred_start + len(original)
        if block_content[preferred_start:end] == original:
            return preferred_start
    return block_content.find(original, search_start)


def _correction_sort_key(block_content, correction):
    preferred_start = correction.get('preferred_start')
    if isinstance(preferred_start, int) and preferred_start >= 0:
        return preferred_start
    fallback = block_content.find(correction.get('original', ''))
    return fallback if fallback >= 0 else len(block_content)


def _filter_corrections_for_block(block_content, corrections):
    """Return only corrections that apply to one paragraph/block of text.

    One entry is emitted per occurrence of the original text in the block so
    that every repeated instance (e.g. missing terminal punctuation appearing
    multiple times) is marked, not just the first one the renderer finds.
    """
    block_corrections = []
    seen = set()

    for corr in corrections:
        original = corr.get('original')
        corrected = corr.get('corrected', original)
        explanation = corr.get('explanation', '')

        if not original or original == corrected or original not in block_content:
            continue

        occurrence_starts = _find_all_indices(block_content, original)
        atomic_edits = _build_atomic_edits(original, corrected)
        if not atomic_edits:
            atomic_edits = [
                {
                    'relative_start': 0,
                    'original': original,
                    'corrected': corrected,
                }
            ]

        for occurrence_start in occurrence_starts:
            for edit in atomic_edits:
                preferred_start = occurrence_start + edit['relative_start']
                entry = {
                    'explanation': explanation,
                    'original': edit['original'],
                    'corrected': edit['corrected'],
                    'preferred_start': preferred_start,
                }
                dedupe_key = (
                    preferred_start,
                    entry['original'],
                    entry['corrected'],
                )
                if dedupe_key in seen:
                    continue
                seen.add(dedupe_key)
                block_corrections.append(entry)

    block_corrections.sort(key=lambda item: _correction_sort_key(block_content, item))
    return block_corrections


def _rewrite_paragraph_preserving_images(para, block_content, block_corrections, config):
    """
    Rewrites the text of a paragraph that contains both text and image runs.
    Drawing/picture nodes are preserved; only text-bearing runs are modified.
    """
    p = para._p

    drawing_nodes = []
    for r_elem in list(p.findall(qn('w:r'))):
        for child_tag in (qn('w:drawing'), qn('w:pict')):
            d = r_elem.find(child_tag)
            if d is not None:
                idx = list(p).index(r_elem)
                drawing_nodes.append((idx, copy.deepcopy(r_elem)))

    for r_elem in list(p.findall(qn('w:r'))):
        p.remove(r_elem)

    def _add_run_elem(text, bold=False, color=None, italic=False, strike=False):
        r = etree.SubElement(p, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        if bold or color or italic or strike:
            rpr = etree.Element(qn('w:rPr'))
            r.insert(0, rpr)
            if bold:
                etree.SubElement(rpr, qn('w:b'))
            if color:
                clr = etree.SubElement(rpr, qn('w:color'))
                clr.set(qn('w:val'), color)
            if italic:
                etree.SubElement(rpr, qn('w:i'))
            if strike:
                etree.SubElement(rpr, qn('w:strike'))
        return r

    block_corrections.sort(key=lambda item: _correction_sort_key(block_content, item))
    last_end = 0
    for corr in block_corrections:
        orig = corr['original']
        start = _resolve_correction_start(
            block_content,
            orig,
            corr.get('preferred_start'),
            last_end,
        )
        if start == -1:
            continue
        end = start + len(orig)

        if block_content[last_end:start]:
            _add_run_elem(block_content[last_end:start])

        corrected_text = _preserve_soft_breaks(orig, corr.get('corrected', orig))
        matcher = difflib.SequenceMatcher(None, orig, corrected_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                _add_run_elem(corrected_text[j1:j2])
            elif tag in ('replace', 'insert'):
                if config.get('highlight_corrections', True):
                    _add_run_elem(corrected_text[j1:j2], bold=True, color='FF0000')
                else:
                    _add_run_elem(corrected_text[j1:j2])
            elif tag == 'delete':
                deleted_marker = _build_deletion_marker(
                    orig[i1:i2],
                    show_marker=config.get('show_deletion_markers', True),
                )
                if deleted_marker:
                    _add_run_elem(deleted_marker, color='FF0000', strike=True)

        if config.get('add_comments', True):
            explanation = corr.get('explanation', '').strip()
            if explanation:
                _add_run_elem(f" [{explanation}]", italic=True, color='0000FF')

        last_end = end

    if block_content[last_end:]:
        _add_run_elem(block_content[last_end:])

    for _idx, r_elem in drawing_nodes:
        p.append(r_elem)


def _append_batch_to_correction_plan(current_batch, config, client, correction_plan, stats):
    """Gets LLM corrections for one batch and appends normalized entries to correction_plan."""
    full_text = "\n".join([b['content'] for b in current_batch])
    corrections, input_tokens, output_tokens, llm_time = get_corrections_from_llm(full_text, config, client)
    stats["total_input_tokens"] += input_tokens
    stats["total_tokens_generated"] += output_tokens
    stats["total_llm_time"] += llm_time

    for item in current_batch:
        block_content = item['content']
        stats["total_text_size"] += len(block_content)
        correction_plan.append({
            'content': block_content,
            'corrections': _filter_corrections_for_block(block_content, corrections)
        })


def _apply_inline_corrections_to_paragraph(para, block_content, block_corrections, config):
    """Apply precomputed corrections to one paragraph in inline mode."""
    if not block_corrections:
        return

    if _paragraph_contains_image(para):
        _rewrite_paragraph_preserving_images(para, block_content, block_corrections, config)
        return

    para.clear()
    block_corrections.sort(key=lambda item: _correction_sort_key(block_content, item))

    last_end = 0
    for corr in block_corrections:
        orig = corr['original']
        start = _resolve_correction_start(
            block_content,
            orig,
            corr.get('preferred_start'),
            last_end,
        )
        if start == -1:
            continue

        end = start + len(orig)
        para.add_run(block_content[last_end:start])

        corrected_text = _preserve_soft_breaks(orig, corr.get('corrected', orig))
        matcher = difflib.SequenceMatcher(None, orig, corrected_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                para.add_run(corrected_text[j1:j2])
            elif tag == 'replace' or tag == 'insert':
                run = para.add_run(corrected_text[j1:j2])
                if config.get('highlight_corrections', True):
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            elif tag == 'delete':
                deleted_marker = _build_deletion_marker(
                    orig[i1:i2],
                    show_marker=config.get('show_deletion_markers', True),
                )
                if deleted_marker:
                    deleted_run = para.add_run(deleted_marker)
                    deleted_run.font.strike = True
                    deleted_run.font.color.rgb = RGBColor(255, 0, 0)

        if config.get('add_comments', True):
            explanation = corr.get('explanation', '').strip()
            if explanation:
                exp_run = para.add_run(f" [{explanation}]")
                exp_run.italic = True
                exp_run.font.color.rgb = RGBColor(0, 0, 255)

        last_end = end

    para.add_run(block_content[last_end:])


def build_correction_plan(input_path, config, client, should_cancel=None):
    """Build one correction plan from LLM output that can drive multiple renderers."""
    def _raise_if_canceled():
        if callable(should_cancel) and should_cancel():
            raise RuntimeError("Canceled by user request")

    _raise_if_canceled()
    print(f"Loading document: {input_path}")
    doc = Document(input_path)
    prompt_key = config.get('active_prompt', DEFAULT_PROMPT_KEY)
    max_input_words = get_prompt_max_input_words(prompt_key, fallback=500)
    print(f"Using prompt context size: {max_input_words} words")

    stats = {
        "total_text_size": 0,
        "total_llm_time": 0,
        "total_input_tokens": 0,
        "total_tokens_generated": 0,
        "hidden_chars_normalized": 0,
    }

    correction_plan = []
    all_paragraphs = _collect_all_paragraphs(doc)

    current_batch = []
    current_word_count = 0

    for para in all_paragraphs:
        _raise_if_canceled()
        text = para.text.strip()
        if not text:
            continue

        if current_word_count + len(text.split()) > max_input_words and current_batch:
            _raise_if_canceled()
            _append_batch_to_correction_plan(current_batch, config, client, correction_plan, stats)

            current_batch = []
            current_word_count = 0

        # Normalize hidden whitespace before sending to LLM (primary prevention point)
        normalized_text, norm_count = _normalize_hidden_whitespace(para.text)
        stats["hidden_chars_normalized"] += norm_count
        
        current_batch.append({'content': normalized_text})
        current_word_count += len(text.split())

    if current_batch:
        _raise_if_canceled()
        _append_batch_to_correction_plan(current_batch, config, client, correction_plan, stats)

    if stats["hidden_chars_normalized"] > 0:
        print(f"[i] Normalized {stats['hidden_chars_normalized']} hidden Unicode whitespace character(s) before LLM analysis.")

    return correction_plan, stats


def apply_inline_correction_plan(input_path, output_path, correction_plan, config):
    """Apply a precomputed correction plan to a DOCX in inline format."""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)

    inserted = _insert_blank_line_before_images(doc)
    if inserted:
        print(f"Inserted {inserted} blank line(s) before image paragraph(s).")

    print("Applying inline corrections from shared correction plan...")

    all_paragraphs = _collect_all_paragraphs(doc)
    paragraphs = [
        {'para': para, 'content': para.text}
        for para in all_paragraphs
        if para.text and para.text.strip()
    ]

    paragraph_cursor = 0
    for item in correction_plan:
        block_corrections = item.get('corrections', [])
        if not block_corrections:
            continue

        matched_paragraph = None
        for idx in range(paragraph_cursor, len(paragraphs)):
            if paragraphs[idx]['content'] == item['content']:
                matched_paragraph = paragraphs[idx]['para']
                paragraph_cursor = idx + 1
                break

        if matched_paragraph is None:
            continue

        _apply_inline_corrections_to_paragraph(
            matched_paragraph,
            item['content'],
            block_corrections,
            config,
        )

    doc.save(output_path)
    print(f"\nSuccessfully saved corrected DOCX to: {output_path}")


def process_docx(input_path, output_path, config, client):
    """Loads a DOCX, corrects text in-place (preserving images), and saves to output_path."""
    correction_plan, stats = build_correction_plan(input_path, config, client)
    apply_inline_correction_plan(input_path, output_path, correction_plan, config)
    return stats


# ---------------------------------------------------------------------------
# Prepend-text output: generate text and insert it at the top of the document
# ---------------------------------------------------------------------------

def build_prepend_plan(input_path, config, client):
    """Collect all document text and send it to the LLM as a single request for plain-text generation."""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)

    stats = {
        "total_text_size": 0,
        "total_llm_time": 0.0,
        "total_input_tokens": 0,
        "total_tokens_generated": 0,
    }

    all_paragraphs = _collect_all_paragraphs(doc)
    full_text = "\n".join(para.text for para in all_paragraphs if para.text.strip())
    stats["total_text_size"] = len(full_text)

    if not full_text.strip():
        print("  [!] Document appears to be empty. No summary generated.")
        return "", stats

    print("Sending document text to LLM for summary generation...")
    generated_text, input_tokens, output_tokens, llm_time = get_text_from_llm(full_text, config, client)

    stats["total_input_tokens"] = input_tokens
    stats["total_tokens_generated"] = output_tokens
    stats["total_llm_time"] = llm_time

    return generated_text.strip(), stats


def apply_prepend_plan(input_path, output_path, prepend_text):
    """Insert a 'Chapter Summary' heading and the generated text at the top of the DOCX."""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)
    body = doc.element.body

    # Add heading and text paragraphs (python-docx appends to end)
    heading = doc.add_heading("Chapter Summary", level=2)
    text_para = doc.add_paragraph(prepend_text)

    # Move them both to the front of the document body
    body.remove(text_para._element)
    body.remove(heading._element)
    body.insert(0, text_para._element)
    body.insert(0, heading._element)

    doc.save(output_path)
    print(f"\nSuccessfully saved document with summary to: {output_path}")


def _extract_document_structure(file_path, max_key_points=15):
    """Extract document structure: title, headings, and key content for course overview."""
    try:
        doc = Document(file_path)
        structure = {
            "title": file_path.stem,
            "headings": [],
            "key_content": [],
        }
        
        # Extract headings and surrounding content
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 3:
                continue
            
            # Capture heading-level content  
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                if level <= 2:  # Keep only top 2 levels of headings
                    structure["headings"].append(text)
            
            # Capture bullet points and numbered lists
            if para.style.name.startswith("List"):
                clean_text = text.lstrip("•-*0123456789.). ").strip()
                if clean_text and len(structure["key_content"]) < max_key_points:
                    structure["key_content"].append(clean_text)
            
            # Also capture first substantive paragraph after headings (intro text)
            elif para.style.name in ("Normal", "Body Text", "Body Text 2"):
                # Look for intro text (paragraph after a heading, relatively short and substantive)
                if 20 < len(text) < 300 and any(w in text.lower() for w in ["learn", "understand", "explore", "examine", "review", "overview", "introduction"]):
                    if len(structure["key_content"]) < max_key_points:
                        # Take first 100 chars of intro text
                        structure["key_content"].append(text[:100] + "..." if len(text) > 100 else text)
        
        return structure
    except Exception as e:
        print(f"  [!] Error extracting structure from {file_path.name}: {e}")
        return None


def build_course_summary_plan(file_paths, config, client):
    """Extract document structures and send structured course outline to LLM for synthesis."""
    print(f"Analyzing {len(file_paths)} document(s) to extract course structure...")
    
    stats = {
        "total_text_size": 0,
        "total_llm_time": 0.0,
        "total_input_tokens": 0,
        "total_tokens_generated": 0,
    }
    
    # Extract structure from all documents
    course_structure = []
    for file_path in file_paths:
        structure = _extract_document_structure(file_path)
        if structure:
            course_structure.append(structure)
            print(f"  Extracted: {structure['title']} ({len(structure['headings'])} sections, {len(structure['key_content'])} topics)")
    
    if not course_structure:
        print("  [!] No valid documents found. No course summary generated.")
        return "", stats
    
    # Build structured summary for LLM
    structured_input = "COURSE STRUCTURE ANALYSIS\n" + "=" * 50 + "\n\n"
    
    for doc_struct in course_structure:
        structured_input += f"MODULE: {doc_struct['title']}\n"
        if doc_struct['headings']:
            structured_input += "Main Topics:\n"
            for heading in doc_struct['headings'][:8]:  # Limit headings
                structured_input += f"  • {heading}\n"
        if doc_struct['key_content']:
            structured_input += "Key Content:\n"
            for point in doc_struct['key_content'][:10]:  # Limit points
                structured_input += f"  • {point}\n"
        structured_input += "\n"
    
    stats["total_text_size"] = len(structured_input)
    
    print("Sending course structure to LLM for synthesis...")
    generated_text, input_tokens, output_tokens, llm_time = get_text_from_llm(structured_input, config, client)
    
    stats["total_input_tokens"] = input_tokens
    stats["total_tokens_generated"] = output_tokens
    stats["total_llm_time"] = llm_time
    
    return generated_text.strip(), stats


def save_course_summary(output_path, course_summary_text):
    """Save the generated course summary as a standalone DOCX file with proper section formatting."""
    from docx.shared import Inches

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    doc.add_heading("Course Summary", level=1)

    # Parse labeled sections [SECTION NAME] from LLM output
    SECTION_LABELS = {
        "[COURSE DESCRIPTION]": "Course Description",
        "[INTENDED AUDIENCE]": "Intended Audience",
        "[COURSE LENGTH]": "Course Length",
        "[LEARNING OBJECTIVES]": "Learning Objectives",
    }

    lines = course_summary_text.split('\n')
    current_section_label = None
    current_prose_lines = []

    def _flush_prose(lines_buf, section_label):
        if lines_buf:
            text = ' '.join(lines_buf).strip()
            if text:
                doc.add_paragraph(text)
            lines_buf.clear()

    for line in lines:
        raw = line.strip()
        upper = raw.upper()

        # Detect section labels — may come without brackets too (LLM sometimes strips them)
        matched_label = None
        for tag, label in SECTION_LABELS.items():
            if upper == tag or upper == tag.strip('[]') or upper.startswith(tag):
                matched_label = label
                break
            # Also match without brackets e.g. "COURSE DESCRIPTION"
            if upper == label.upper() or upper.startswith(label.upper() + ':'):
                matched_label = label
                break

        if matched_label:
            _flush_prose(current_prose_lines, current_section_label)
            current_section_label = matched_label
            doc.add_heading(matched_label, level=2)
            continue

        if not raw:
            _flush_prose(current_prose_lines, current_section_label)
            continue

        if raw.startswith(('•', '-', '*')) and current_section_label == "Learning Objectives":
            _flush_prose(current_prose_lines, current_section_label)
            clean_text = raw.lstrip('•-* ').strip()
            if clean_text:
                doc.add_paragraph(clean_text, style='List Bullet')
        else:
            current_prose_lines.append(raw)

    _flush_prose(current_prose_lines, current_section_label)

    doc.save(output_path)
    print(f"\nSuccessfully saved course summary to: {output_path}")


# ---------------------------------------------------------------------------
# Hybrid output: inline red text changes + Word XML comment explanations
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


def _apply_hybrid_corrections_to_paragraph(para, block_content, block_corrections, config, comment_collector):
    """Inline red text changes with explanations stored as Word XML comments instead of inline text."""
    if not block_corrections:
        return

    if _paragraph_contains_image(para):
        # Apply inline changes to image paragraphs but suppress inline explanation text.
        hybrid_config = dict(config)
        hybrid_config['add_comments'] = False
        _rewrite_paragraph_preserving_images(para, block_content, block_corrections, hybrid_config)
        # Gather all explanations and attach them as a single comment spanning the
        # paragraph, since individual run anchoring is not feasible after the rewrite.
        explanations = [
            corr.get('explanation', '').strip()
            for corr in block_corrections
            if corr.get('explanation', '').strip()
        ]
        if explanations:
            p = para._p
            comment_id = len(comment_collector)
            comment_start = etree.Element(qn('w:commentRangeStart'))
            comment_start.set(qn('w:id'), str(comment_id))
            p.insert(0, comment_start)
            comment_end = etree.Element(qn('w:commentRangeEnd'))
            comment_end.set(qn('w:id'), str(comment_id))
            p.append(comment_end)
            ref_run = etree.SubElement(p, qn('w:r'))
            ref_rpr = etree.SubElement(ref_run, qn('w:rPr'))
            ref_style = etree.SubElement(ref_rpr, qn('w:rStyle'))
            ref_style.set(qn('w:val'), 'CommentReference')
            comment_ref = etree.SubElement(ref_run, qn('w:commentReference'))
            comment_ref.set(qn('w:id'), str(comment_id))
            comment_collector.append((comment_id, ' | '.join(explanations)))
        return

    para.clear()
    p = para._p
    block_corrections.sort(key=lambda item: _correction_sort_key(block_content, item))

    last_end = 0
    for corr in block_corrections:
        orig = corr['original']
        start = _resolve_correction_start(
            block_content,
            orig,
            corr.get('preferred_start'),
            last_end,
        )
        if start == -1:
            continue
        end = start + len(orig)

        prefix = block_content[last_end:start]
        if prefix:
            para.add_run(prefix)

        explanation = corr.get('explanation', '').strip()
        comment_id = None
        if explanation:
            comment_id = len(comment_collector)
            comment_start = etree.Element(qn('w:commentRangeStart'))
            comment_start.set(qn('w:id'), str(comment_id))
            p.append(comment_start)

        corrected_text = _preserve_soft_breaks(orig, corr.get('corrected', orig))
        matcher = difflib.SequenceMatcher(None, orig, corrected_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                para.add_run(corrected_text[j1:j2])
            elif tag in ('replace', 'insert'):
                run = para.add_run(corrected_text[j1:j2])
                if config.get('highlight_corrections', True):
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            elif tag == 'delete':
                deleted_marker = _build_deletion_marker(
                    orig[i1:i2],
                    show_marker=config.get('show_deletion_markers', True),
                )
                if deleted_marker:
                    deleted_run = para.add_run(deleted_marker)
                    deleted_run.font.strike = True
                    deleted_run.font.color.rgb = RGBColor(255, 0, 0)

        if explanation:
            comment_end = etree.Element(qn('w:commentRangeEnd'))
            comment_end.set(qn('w:id'), str(comment_id))
            p.append(comment_end)

            ref_run = etree.SubElement(p, qn('w:r'))
            ref_run_pr = etree.SubElement(ref_run, qn('w:rPr'))
            ref_style = etree.SubElement(ref_run_pr, qn('w:rStyle'))
            ref_style.set(qn('w:val'), 'CommentReference')
            comment_ref = etree.SubElement(ref_run, qn('w:commentReference'))
            comment_ref.set(qn('w:id'), str(comment_id))

            comment_collector.append((comment_id, explanation))

        last_end = end

    suffix = block_content[last_end:]
    if suffix:
        para.add_run(suffix)


def _build_comments_xml(comment_items):
    """Return bytes for word/comments.xml from a list of (comment_id, text) pairs."""
    root = etree.Element(f'{{{_W_NS}}}comments', nsmap={'w': _W_NS})
    for comment_id, comment_text in comment_items:
        comment = etree.SubElement(root, f'{{{_W_NS}}}comment')
        comment.set(f'{{{_W_NS}}}id', str(comment_id))
        comment.set(f'{{{_W_NS}}}author', 'AI Reviewer')
        comment.set(f'{{{_W_NS}}}date', '2026-01-01T00:00:00Z')
        comment.set(f'{{{_W_NS}}}initials', 'AI')

        paragraph = etree.SubElement(comment, f'{{{_W_NS}}}p')
        run = etree.SubElement(paragraph, f'{{{_W_NS}}}r')
        text = etree.SubElement(run, f'{{{_W_NS}}}t')
        text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        text.text = comment_text

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _inject_comments_into_docx(docx_path, comment_items):
    """Post-process a saved docx zip to add word/comments.xml and required relationships."""
    import re as _re
    import zipfile as _zipfile

    comments_xml = _build_comments_xml(comment_items)

    with _zipfile.ZipFile(docx_path, 'r') as source_zip:
        file_map = {name: source_zip.read(name) for name in source_zip.namelist()}

    rels_key = 'word/_rels/document.xml.rels'
    if rels_key in file_map:
        rels_text = file_map[rels_key].decode('utf-8')
        if _COMMENTS_REL_TYPE not in rels_text:
            existing_ids = set(_re.findall(r'Id="([^"]+)"', rels_text))
            rel_id = 'rIdComments'
            counter = 0
            while rel_id in existing_ids:
                counter += 1
                rel_id = f'rIdComments{counter}'

            relation = (
                f'<Relationship Id="{rel_id}" '
                f'Type="{_COMMENTS_REL_TYPE}" '
                'Target="comments.xml"/>'
            )
            rels_text = rels_text.replace('</Relationships>', f'{relation}</Relationships>')
            file_map[rels_key] = rels_text.encode('utf-8')

    content_types_key = '[Content_Types].xml'
    if content_types_key in file_map:
        content_types = file_map[content_types_key].decode('utf-8')
        if 'comments.xml' not in content_types:
            override = (
                '<Override PartName="/word/comments.xml" '
                f'ContentType="{_COMMENTS_CONTENT_TYPE}"/>'
            )
            content_types = content_types.replace('</Types>', f'{override}</Types>')
            file_map[content_types_key] = content_types.encode('utf-8')

    file_map['word/comments.xml'] = comments_xml

    with _zipfile.ZipFile(docx_path, 'w', _zipfile.ZIP_DEFLATED) as target_zip:
        for name, content in file_map.items():
            target_zip.writestr(name, content)


def apply_hybrid_correction_plan(input_path, output_path, correction_plan, config):
    """Apply a precomputed correction plan to a DOCX in hybrid format."""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)

    inserted = _insert_blank_line_before_images(doc)
    if inserted:
        print(f"Inserted {inserted} blank line(s) before image paragraph(s).")

    print("Applying hybrid corrections from shared correction plan...")

    all_paragraphs = _collect_all_paragraphs(doc)
    paragraphs = [
        {'para': para, 'content': para.text}
        for para in all_paragraphs
        if para.text and para.text.strip()
    ]

    comment_collector = []
    paragraph_cursor = 0

    for item in correction_plan:
        block_corrections = item.get('corrections', [])
        if not block_corrections:
            continue

        matched_paragraph = None
        for idx in range(paragraph_cursor, len(paragraphs)):
            if paragraphs[idx]['content'] == item['content']:
                matched_paragraph = paragraphs[idx]['para']
                paragraph_cursor = idx + 1
                break

        if matched_paragraph is None:
            continue

        _apply_hybrid_corrections_to_paragraph(
            matched_paragraph,
            item['content'],
            block_corrections,
            config,
            comment_collector,
        )

    doc.save(output_path)

    if comment_collector:
        _inject_comments_into_docx(output_path, comment_collector)

    print(f"\nSuccessfully saved hybrid DOCX to: {output_path}")
