"""Deterministic run summary state and DOCX rendering."""

import ctypes
import json
import os
import re
import tempfile
import time
from collections import Counter
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path

from docx import Document

SCHEMA_VERSION = 1
MAX_RUN_HISTORY = 200
CATEGORY_ORDER = [
    "spelling",
    "grammar",
    "punctuation",
    "capitalization",
    "terminology",
    "style",
    "formatting",
    "other",
]

CATEGORY_LABELS = {
    "spelling": "Spelling",
    "grammar": "Grammar",
    "punctuation": "Punctuation",
    "capitalization": "Capitalization",
    "terminology": "Terminology",
    "style": "Style",
    "formatting": "Formatting",
    "other": "Other",
}

CATEGORY_KEYWORDS = {
    "spelling": ["spelling", "misspell", "typo", "misspelled"],
    "grammar": ["grammar", "agreement", "tense", "article", "subject-verb"],
    "punctuation": ["punctuation", "comma", "period", "apostrophe", "quote", "semicolon", "colon"],
    "capitalization": ["capitalization", "uppercase", "lowercase", "capitalized"],
    "terminology": ["terminology", "term", "product name", "brand", "naming"],
    "style": ["style", "clarity", "readability", "concise", "wording", "tone"],
    "formatting": ["formatting", "whitespace", "line break", "spacing", "bullet", "numbering"],
}


def _now_iso() -> str:
    return datetime.now().isoformat()


def _empty_state() -> dict:
    now = _now_iso()
    return {
        "schemaVersion": SCHEMA_VERSION,
        "createdAt": now,
        "updatedAt": now,
        "runs": [],
    }


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip()).lower()


def categorize_correction(correction: dict) -> str:
    """Assign a deterministic category to one correction entry."""
    explanation = _normalize_text(correction.get("explanation", ""))
    original = correction.get("original", "") or ""
    corrected = correction.get("corrected", original) or ""

    for category, keywords in CATEGORY_KEYWORDS.items():
        if any(keyword in explanation for keyword in keywords):
            return category

    if original != corrected and original.lower() == corrected.lower():
        return "capitalization"

    punctuation_pattern = r"[\w\s]"
    original_marks = re.sub(punctuation_pattern, "", original)
    corrected_marks = re.sub(punctuation_pattern, "", corrected)
    if original_marks != corrected_marks and _normalize_text(original) == _normalize_text(corrected):
        return "punctuation"

    original_words = [word for word in re.split(r"\s+", original.strip()) if word]
    corrected_words = [word for word in re.split(r"\s+", corrected.strip()) if word]
    if len(original_words) == 1 and len(corrected_words) == 1:
        return "spelling"

    return "other"


def summarize_correction_plan(correction_plan: list[dict]) -> dict:
    """Create correction counts per category from one correction plan."""
    category_counts = Counter({category: 0 for category in CATEGORY_ORDER})
    correction_count = 0

    for block in correction_plan:
        for correction in block.get("corrections", []):
            correction_count += 1
            category = categorize_correction(correction)
            category_counts[category] += 1

    return {
        "correctionCount": correction_count,
        "categoryCounts": {category: category_counts.get(category, 0) for category in CATEGORY_ORDER},
    }


def load_state(path: Path) -> dict:
    if not path.exists():
        return _empty_state()

    payload = _read_json_with_retry(path)
    if payload is None:
        return _empty_state()

    if not isinstance(payload, dict) or not isinstance(payload.get("runs"), list):
        return _empty_state()

    payload["schemaVersion"] = SCHEMA_VERSION
    payload.setdefault("createdAt", _now_iso())
    payload["updatedAt"] = _now_iso()
    return payload


def save_state(path: Path, state: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    state["updatedAt"] = _now_iso()
    _write_json_atomic_with_retry(path, state)
    _mark_hidden_on_windows(path)


def _read_json_with_retry(path: Path, attempts: int = 5, delay_seconds: float = 0.2) -> dict | None:
    last_error: Exception | None = None
    for attempt in range(attempts):
        try:
            with open(path, "r", encoding="utf-8") as file_handle:
                return json.load(file_handle)
        except FileNotFoundError:
            return None
        except PermissionError as exc:
            last_error = exc
            if attempt < attempts - 1:
                time.sleep(delay_seconds)
                continue
            raise
        except Exception as exc:
            last_error = exc
            break

    if last_error is not None and isinstance(last_error, PermissionError):
        raise last_error
    return None


def _write_json_atomic_with_retry(path: Path, payload: dict, attempts: int = 6, delay_seconds: float = 0.25) -> None:
    last_error: PermissionError | None = None
    for attempt in range(attempts):
        temp_fd = None
        temp_path = None
        try:
            temp_fd, temp_name = tempfile.mkstemp(prefix="summary_report_state_", suffix=".tmp", dir=str(path.parent))
            temp_path = Path(temp_name)
            with os.fdopen(temp_fd, "w", encoding="utf-8") as file_handle:
                temp_fd = None
                json.dump(payload, file_handle, indent=2)

            os.replace(temp_path, path)
            return
        except PermissionError as exc:
            last_error = exc
            if attempt < attempts - 1:
                time.sleep(delay_seconds)
                continue
            raise
        finally:
            if temp_fd is not None:
                os.close(temp_fd)
            if temp_path is not None and temp_path.exists():
                try:
                    temp_path.unlink()
                except Exception:
                    pass

    if last_error is not None:
        raise last_error


@contextmanager
def _summary_update_lock(lock_path: Path, wait_seconds: float = 10.0, poll_seconds: float = 0.1):
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    start = time.time()
    while True:
        try:
            fd = os.open(str(lock_path), os.O_CREAT | os.O_EXCL | os.O_RDWR)
            break
        except FileExistsError:
            if (time.time() - start) >= wait_seconds:
                raise TimeoutError(f"Timed out waiting for summary report lock: {lock_path}")
            time.sleep(poll_seconds)

    try:
        with os.fdopen(fd, "w", encoding="utf-8") as lock_file:
            lock_file.write(f"pid={os.getpid()}\n")
            lock_file.write(f"timestamp={_now_iso()}\n")
            lock_file.flush()
        yield
    finally:
        try:
            lock_path.unlink()
        except FileNotFoundError:
            pass


def _mark_hidden_on_windows(path: Path) -> None:
    if os.name != "nt":
        return
    try:
        file_attr_hidden = 0x02
        file_attr_normal = 0x80
        get_file_attributes = ctypes.windll.kernel32.GetFileAttributesW
        set_file_attributes = ctypes.windll.kernel32.SetFileAttributesW
        get_file_attributes.restype = ctypes.c_uint32

        attrs = get_file_attributes(str(path))
        if attrs == 0xFFFFFFFF:
            set_file_attributes(str(path), file_attr_hidden)
            return

        if attrs & file_attr_hidden:
            return

        new_attrs = attrs | file_attr_hidden
        if new_attrs == 0:
            new_attrs = file_attr_normal | file_attr_hidden
        set_file_attributes(str(path), new_attrs)
    except Exception:
        return


def append_run(state: dict, run_record: dict) -> dict:
    runs = list(state.get("runs", []))
    runs.append(run_record)
    if len(runs) > MAX_RUN_HISTORY:
        runs = runs[-MAX_RUN_HISTORY:]
    state["runs"] = runs
    return state


def _compute_global_totals(runs: list[dict]) -> dict:
    category_counts = Counter({category: 0 for category in CATEGORY_ORDER})
    total_files = 0
    total_corrections = 0

    for run in runs:
        total_files += int(run.get("fileCount", 0))
        total_corrections += int(run.get("correctionCount", 0))
        for category, count in run.get("categoryCounts", {}).items():
            if category in category_counts:
                category_counts[category] += int(count or 0)

    return {
        "totalRuns": len(runs),
        "completedRuns": sum(1 for run in runs if run.get("status") == "completed"),
        "canceledRuns": sum(1 for run in runs if run.get("status") == "canceled"),
        "failedRuns": sum(1 for run in runs if run.get("status") == "failed"),
        "totalFiles": total_files,
        "totalCorrections": total_corrections,
        "categoryCounts": {category: category_counts.get(category, 0) for category in CATEGORY_ORDER},
    }


def _compute_document_totals(runs: list[dict]) -> dict[str, dict]:
    by_document: dict[str, dict] = {}

    for run in runs:
        for file_entry in run.get("files", []):
            name = file_entry.get("name", "unknown")
            current = by_document.setdefault(
                name,
                {
                    "runs": 0,
                    "corrections": 0,
                    "categoryCounts": Counter({category: 0 for category in CATEGORY_ORDER}),
                },
            )
            current["runs"] += 1
            current["corrections"] += int(file_entry.get("correctionCount", 0))
            for category, count in file_entry.get("categoryCounts", {}).items():
                if category in current["categoryCounts"]:
                    current["categoryCounts"][category] += int(count or 0)

    return by_document


def render_docx(state: dict, docx_path: Path) -> None:
    runs = list(state.get("runs", []))
    totals = _compute_global_totals(runs)
    document_totals = _compute_document_totals(runs)

    doc = Document()
    doc.add_heading("Summary Report", level=1)
    doc.add_paragraph(f"Updated: {state.get('updatedAt', '')}")
    doc.add_paragraph(f"Runs tracked: {totals['totalRuns']} (completed {totals['completedRuns']}, canceled {totals['canceledRuns']}, failed {totals['failedRuns']})")
    doc.add_paragraph(f"Files processed: {totals['totalFiles']}")
    doc.add_paragraph(f"Corrections applied: {totals['totalCorrections']}")

    doc.add_heading("Category Totals", level=2)
    category_table = doc.add_table(rows=1, cols=2)
    category_table.style = "Light Grid"
    category_table.rows[0].cells[0].text = "Category"
    category_table.rows[0].cells[1].text = "Corrections"
    for category in CATEGORY_ORDER:
        row = category_table.add_row().cells
        row[0].text = CATEGORY_LABELS.get(category, category.title())
        row[1].text = str(totals["categoryCounts"].get(category, 0))

    doc.add_heading("Recent Runs", level=2)
    run_table = doc.add_table(rows=1, cols=7)
    run_table.style = "Light Grid"
    headers = ["Timestamp", "Status", "Prompt", "Model", "Files", "Corrections", "LLM Time (s)"]
    for index, header in enumerate(headers):
        run_table.rows[0].cells[index].text = header

    for run in reversed(runs[-15:]):
        row = run_table.add_row().cells
        row[0].text = str(run.get("timestamp", ""))
        row[1].text = str(run.get("status", ""))
        row[2].text = str(run.get("promptKey", ""))
        row[3].text = str(run.get("model", ""))
        row[4].text = str(run.get("fileCount", 0))
        row[5].text = str(run.get("correctionCount", 0))
        row[6].text = f"{float(run.get('totalLlmTime', 0.0)):.2f}"

    doc.add_heading("Per-Document Totals", level=2)
    doc_table = doc.add_table(rows=1, cols=4)
    doc_table.style = "Light Grid"
    doc_table.rows[0].cells[0].text = "Document"
    doc_table.rows[0].cells[1].text = "Runs"
    doc_table.rows[0].cells[2].text = "Corrections"
    doc_table.rows[0].cells[3].text = "Top Category"

    sorted_docs = sorted(document_totals.items(), key=lambda item: item[1]["corrections"], reverse=True)
    for name, totals_by_doc in sorted_docs:
        row = doc_table.add_row().cells
        row[0].text = name
        row[1].text = str(totals_by_doc["runs"])
        row[2].text = str(totals_by_doc["corrections"])
        top_category = max(CATEGORY_ORDER, key=lambda category: totals_by_doc["categoryCounts"].get(category, 0))
        row[3].text = CATEGORY_LABELS.get(top_category, top_category.title())

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def update_summary_report(output_dir: Path, run_record: dict) -> dict:
    """Merge one run into summary state and regenerate DOCX report."""
    state_path = output_dir / "summary_report_state.json"
    docx_path = output_dir / "summary_report.docx"

    lock_path = output_dir / ".summary_report.lock"
    with _summary_update_lock(lock_path):
        state = load_state(state_path)
        append_run(state, run_record)
        save_state(state_path, state)
        render_docx(state, docx_path)

    return {
        "statePath": state_path,
        "reportPath": docx_path,
    }
