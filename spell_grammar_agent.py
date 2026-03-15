import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from openai import OpenAI
from utils import load_config
from llm_service import get_corrections_from_llm

def process_docx(file_path, config, client):
    doc = Document(file_path)
    new_doc = Document()

    for para in doc.paragraphs:
        original_text = para.text
        if not original_text.strip():
            new_doc.add_paragraph()
            continue

        corrections = get_corrections_from_llm(original_text, config, client)
        new_para = new_doc.add_paragraph()

        if not corrections:
            new_para.add_run(original_text)
            continue

        # Sort corrections by position
        corrections.sort(key=lambda x: original_text.find(x['original']))

        last_end = 0
        for corr in corrections:
            orig = corr['original']
            corrected = corr['corrected']
            explanation = corr['explanation']
            start = original_text.find(orig, last_end)
            if start == -1:
                continue
            end = start + len(orig)
            # Add text before
            new_para.add_run(original_text[last_end:start])
            # Add corrected with highlight
            run = new_para.add_run(corrected)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            # Add explanation inline
            exp_run = new_para.add_run(f" ({explanation})")
            exp_run.italic = True
            exp_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for explanation
            last_end = end
        # Add remaining text
        new_para.add_run(original_text[last_end:])

    # Save as .docx
    output_path = Path(config['output_dir']) / f"{file_path.stem}_corrected.docx"
    new_doc.save(output_path)
    print(f"Processed {file_path.name} -> {output_path.name}")

def main():
    config = load_config()
    workspace_dir = Path(__file__).parent
    input_dir = workspace_dir / config['input_dir']
    config['output_dir'] = str(workspace_dir / config['output_dir'])

    if not input_dir.exists():
        print(f"Input directory {input_dir} does not exist. Please create it and add .docx files.")
        return

    # Initialize LLM client
    client = OpenAI(base_url="http://localhost:11434/v1", api_key="not-needed")

    for file_path in input_dir.glob("*.docx"):
        if "_corrected" in file_path.name:
            continue
        try:
            process_docx(file_path, config, client)
        except Exception as e:
            print(f"Error processing {file_path}: {e}")

if __name__ == "__main__":
    main()