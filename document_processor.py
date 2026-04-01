import difflib
import copy
from lxml import etree
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from llm_service import get_corrections_from_llm


def _paragraph_contains_image(paragraph):
    """Returns True if the paragraph contains an inline/floating image."""
    p = paragraph._p
    return (
        p.find(f".//{qn('w:drawing')}") is not None
        or p.find(f".//{qn('w:pict')}") is not None
    )


def _insert_blank_line_before_images(doc):
    """Ensures a blank paragraph exists before and after every image paragraph."""
    inserted_count = 0
    for para in list(doc.paragraphs):
        if not _paragraph_contains_image(para):
            continue

        prev = para._p.getprevious()
        prev_is_blank_para = (
            prev is not None
            and prev.tag.endswith("}p")
            and "".join(prev.itertext()).strip() == ""
        )

        if not prev_is_blank_para:
            para.insert_paragraph_before("")
            inserted_count += 1

        nxt = para._p.getnext()
        next_is_blank_para = (
            nxt is not None
            and nxt.tag.endswith("}p")
            and "".join(nxt.itertext()).strip() == ""
        )

        if not next_is_blank_para:
            # Insert an empty paragraph after the image paragraph.
            para._p.addnext(copy.deepcopy(para._p))
            new_next = para._p.getnext()
            for child in list(new_next):
                new_next.remove(child)
            inserted_count += 1

    return inserted_count


def _preserve_soft_breaks(original_text, corrected_text):
    """Reinsert original soft line breaks (Shift+Enter) if the correction dropped them."""
    break_char = "\n"
    original_breaks = original_text.count(break_char)
    if original_breaks == 0:
        return corrected_text
    if corrected_text.count(break_char) >= original_breaks:
        return corrected_text

    parts = original_text.split(break_char)
    if len(parts) <= 1:
        return corrected_text

    total_original_len = sum(len(p) for p in parts)
    if total_original_len <= 0:
        return corrected_text

    corrected_len = len(corrected_text)
    # Allocate corrected text across original line proportions, then join with soft breaks.
    target_lengths = [round((len(p) / total_original_len) * corrected_len) for p in parts]
    diff = corrected_len - sum(target_lengths)
    if target_lengths:
        target_lengths[-1] += diff

    chunks = []
    cursor = 0
    for i, target in enumerate(target_lengths):
        remaining = corrected_len - cursor
        if i == len(target_lengths) - 1:
            chunk_len = remaining
        else:
            chunk_len = max(0, min(int(target), remaining))
        chunks.append(corrected_text[cursor:cursor + chunk_len])
        cursor += chunk_len

    return break_char.join(chunks)


def _build_deletion_marker(deleted_text):
    """Builds a visible marker for deleted text so removals are obvious in output."""
    if not deleted_text:
        return ""
    visible = deleted_text.replace("\n", r"\n")
    return f"[-{visible}-]"


def _collect_all_paragraphs(doc):
    """Collect paragraphs from the document body and table cells."""
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    return all_paragraphs


def _filter_corrections_for_block(block_content, corrections):
    """Return corrections that apply to one paragraph/block of text."""
    block_corrections = []
    for corr in corrections:
        if corr.get('original') and corr['original'] in block_content:
            if corr.get('original') != corr.get('corrected'):
                block_corrections.append(corr)
    return block_corrections

def _rewrite_paragraph_preserving_images(para, block_content, block_corrections, config):
    """
    Rewrites the text of a paragraph that contains both text and image runs.
    Drawing/picture nodes are preserved; only text-bearing runs are modified.
    """
    p = para._p

    # Collect drawing elements that must be preserved, keyed by their position
    # We'll detach them, rebuild text runs, then reattach drawings in order.
    drawing_nodes = []
    for r_elem in list(p.findall(qn('w:r'))):
        for child_tag in (qn('w:drawing'), qn('w:pict')):
            d = r_elem.find(child_tag)
            if d is not None:
                # Record position index among all children of <w:p>
                idx = list(p).index(r_elem)
                drawing_nodes.append((idx, copy.deepcopy(r_elem)))

    # Remove all existing <w:r> children (we'll rebuild text ones below)
    for r_elem in list(p.findall(qn('w:r'))):
        p.remove(r_elem)

    # --- Build new text runs (same logic as the normal path) ---
    def _add_run_elem(text, bold=False, color=None, italic=False, strike=False):
        r = etree.SubElement(p, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        if bold or color or italic or strike:
            rpr = etree.Element(qn('w:rPr'))
            r.insert(0, rpr)
            if bold:
                etree.SubElement(rpr, qn('w:b'))
            if color:
                clr = etree.SubElement(rpr, qn('w:color'))
                clr.set(qn('w:val'), color)
            if italic:
                etree.SubElement(rpr, qn('w:i'))
            if strike:
                etree.SubElement(rpr, qn('w:strike'))
        return r

    block_corrections.sort(key=lambda x: block_content.find(x['original']))
    last_end = 0
    for corr in block_corrections:
        orig = corr['original']
        start = block_content.find(orig, last_end)
        if start == -1:
            continue
        end = start + len(orig)

        if block_content[last_end:start]:
            _add_run_elem(block_content[last_end:start])

        corrected_text = _preserve_soft_breaks(orig, corr.get('corrected', orig))
        matcher = difflib.SequenceMatcher(None, orig, corrected_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                _add_run_elem(corrected_text[j1:j2])
            elif tag in ('replace', 'insert'):
                if config.get('highlight_corrections', True):
                    _add_run_elem(corrected_text[j1:j2], bold=True, color='FF0000')
                else:
                    _add_run_elem(corrected_text[j1:j2])
            elif tag == 'delete':
                deleted_marker = _build_deletion_marker(orig[i1:i2])
                if deleted_marker:
                    _add_run_elem(deleted_marker, color='FF0000', strike=True)

        if config.get('add_comments', True):
            explanation = corr.get('explanation', '').strip()
            if explanation:
                _add_run_elem(f" [{explanation}]", italic=True, color='0000FF')

        last_end = end

    if block_content[last_end:]:
        _add_run_elem(block_content[last_end:])

    # Re-insert drawing runs at their original positions
    # Since we cleared all runs, just append them at the end to preserve their content
    for _idx, r_elem in drawing_nodes:
        p.append(r_elem)


def _apply_inline_corrections_to_paragraph(para, block_content, block_corrections, config):
    """Apply precomputed corrections to one python-docx paragraph."""
    if not block_corrections:
        return

    if _paragraph_contains_image(para):
        _rewrite_paragraph_preserving_images(para, block_content, block_corrections, config)
        return

    para.clear()
    block_corrections.sort(key=lambda x: block_content.find(x['original']))

    last_end = 0
    for corr in block_corrections:
        orig = corr['original']
        start = block_content.find(orig, last_end)
        if start == -1:
            continue

        end = start + len(orig)
        para.add_run(block_content[last_end:start])

        corrected_text = _preserve_soft_breaks(orig, corr.get('corrected', orig))
        matcher = difflib.SequenceMatcher(None, orig, corrected_text)
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                para.add_run(corrected_text[j1:j2])
            elif tag in ('replace', 'insert'):
                run = para.add_run(corrected_text[j1:j2])
                if config.get('highlight_corrections', True):
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            elif tag == 'delete':
                deleted_marker = _build_deletion_marker(orig[i1:i2])
                if deleted_marker:
                    deleted_run = para.add_run(deleted_marker)
                    deleted_run.font.strike = True
                    deleted_run.font.color.rgb = RGBColor(255, 0, 0)

        if config.get('add_comments', True):
            explanation = corr.get('explanation', '').strip()
            if explanation:
                exp_run = para.add_run(f" [{explanation}]")
                exp_run.italic = True
                exp_run.font.color.rgb = RGBColor(0, 0, 255)

        last_end = end

    para.add_run(block_content[last_end:])


def _plan_corrections_for_batch(batch_items, config, client, stats):
    """Generate one LLM response for a batch and convert it into per-paragraph corrections."""
    full_text = "\n".join([b['content'] for b in batch_items])

    corrections, tokens, llm_time = get_corrections_from_llm(full_text, config, client)
    stats["total_tokens_generated"] += tokens
    stats["total_llm_time"] += llm_time

    planned_items = []
    for item in batch_items:
        block_content = item['content']
        stats["total_text_size"] += len(block_content)
        planned_items.append({
            'position': item['position'],
            'content': block_content,
            'corrections': _filter_corrections_for_block(block_content, corrections),
        })

    return planned_items


def build_correction_plan(input_path, config, client):
    """Build a reusable correction plan for a DOCX with a single set of LLM calls."""
    print(f"Loading document: {input_path}")
    doc = Document(input_path)

    inserted = _insert_blank_line_before_images(doc)
    if inserted:
        print(f"Inserted {inserted} blank line(s) before image paragraph(s).")

    stats = {
        "total_text_size": 0,
        "total_llm_time": 0,
        "total_tokens_generated": 0,
    }
    print("Preparing correction plan...")

    all_paragraphs = _collect_all_paragraphs(doc)
    current_batch = []
    current_word_count = 0
    correction_plan = []

    for position, para in enumerate(all_paragraphs):
        text = para.text.strip()
        if not text:
            continue

        item = {'position': position, 'content': para.text}
        word_count = len(text.split())
        if current_word_count + word_count > 500 and current_batch:
            correction_plan.extend(_plan_corrections_for_batch(current_batch, config, client, stats))
            current_batch = []
            current_word_count = 0

        current_batch.append(item)
        current_word_count += word_count

    if current_batch:
        correction_plan.extend(_plan_corrections_for_batch(current_batch, config, client, stats))

    return correction_plan, stats


def apply_inline_correction_plan(input_path, output_path, correction_plan, config):
    """Apply a precomputed correction plan to a DOCX and save inline output."""
    doc = Document(input_path)
    _insert_blank_line_before_images(doc)
    all_paragraphs = _collect_all_paragraphs(doc)

    for item in correction_plan:
        position = item['position']
        if position >= len(all_paragraphs):
            continue
        _apply_inline_corrections_to_paragraph(
            all_paragraphs[position],
            item['content'],
            item['corrections'],
            config,
        )

    doc.save(output_path)
    print(f"\nSuccessfully saved corrected DOCX to: {output_path}")

def process_docx(input_path, output_path, config, client):
    """Loads a DOCX, corrects text in-place (preserving images), and saves to output_path."""
    correction_plan, stats = build_correction_plan(input_path, config, client)
    apply_inline_correction_plan(input_path, output_path, correction_plan, config)
    return stats