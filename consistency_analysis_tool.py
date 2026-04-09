import argparse
import json
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from openai import AzureOpenAI, OpenAI

from utils import load_config

OLLAMA_PROVIDER = "ollama"
AZURE_PROVIDER = "azure_openai"
AZURE_AI_FOUNDRY_PROVIDER = "azure_ai_foundry"


def normalize_provider(provider: str) -> str:
    provider = (provider or "").strip().lower()
    if provider in {"azure", "azure_openai", "github"}:
        return AZURE_PROVIDER
    if provider in {"azure_ai_foundry", "foundry"}:
        return AZURE_AI_FOUNDRY_PROVIDER
    return OLLAMA_PROVIDER


def get_azure_settings(config: dict) -> dict:
    return {
        "api_key": os.getenv("AZURE_OPENAI_API_KEY"),
        "endpoint": os.getenv("AZURE_OPENAI_ENDPOINT"),
        "api_version": os.getenv("AZURE_OPENAI_API_VERSION") or config.get("azure_api_version", "2024-10-21"),
        "deployment_name": config.get("azure_deployment_name", "").strip(),
    }


def get_azure_ai_foundry_settings(config: dict) -> dict:
    return {
        "api_key": os.getenv("AZURE_AI_FOUNDRY_API_KEY"),
        "endpoint": os.getenv("AZURE_AI_FOUNDRY_ENDPOINT"),
        "model_name": config.get("azure_ai_foundry_model_name", "").strip(),
    }


def resolve_model_name(config: dict) -> str:
    provider = normalize_provider(config.get("llm_provider", OLLAMA_PROVIDER))
    if provider == AZURE_PROVIDER:
        deployment_name = config.get("azure_deployment_name", "").strip()
        if deployment_name:
            return deployment_name
    if provider == AZURE_AI_FOUNDRY_PROVIDER:
        model_name = config.get("azure_ai_foundry_model_name", "").strip()
        if model_name:
            return model_name
    return config.get("llm_model", "")


def create_client(provider: str, config: dict):
    normalized_provider = normalize_provider(provider)
    if normalized_provider == AZURE_PROVIDER:
        azure_settings = get_azure_settings(config)
        if not azure_settings["api_key"]:
            raise RuntimeError("Missing AZURE_OPENAI_API_KEY environment variable.")
        if not azure_settings["endpoint"]:
            raise RuntimeError("Missing AZURE_OPENAI_ENDPOINT environment variable.")
        return AzureOpenAI(
            api_key=azure_settings["api_key"],
            azure_endpoint=azure_settings["endpoint"],
            api_version=azure_settings["api_version"],
        )
    if normalized_provider == AZURE_AI_FOUNDRY_PROVIDER:
        settings = get_azure_ai_foundry_settings(config)
        if not settings["api_key"]:
            raise RuntimeError("Missing AZURE_AI_FOUNDRY_API_KEY environment variable.")
        if not settings["endpoint"]:
            raise RuntimeError("Missing AZURE_AI_FOUNDRY_ENDPOINT environment variable.")
        return OpenAI(api_key=settings["api_key"], base_url=settings["endpoint"].rstrip("/"))
    return OpenAI(base_url="http://localhost:11434/v1", api_key="ollama")


def build_llm_prompt(metadata: dict) -> str:
    doc_summaries = []
    for doc in metadata.get("documents", []):
        summary = {
            "relative_path": doc.get("relative_path"),
            "status": doc.get("status"),
            "word_count": doc.get("word_count"),
            "headings": doc.get("headings", [])[:15],
            "sample_paragraphs": doc.get("sample_paragraphs", [])[:6],
            "top_keywords": doc.get("top_keywords", [])[:15],
            "product_name_candidates": doc.get("product_name_candidates", [])[:20],
        }
        doc_summaries.append(summary)

    compact_metadata = {
        "generated_at": metadata.get("generated_at"),
        "input_folder": metadata.get("input_folder"),
        "document_count": metadata.get("document_count", 0),
        "global_top_keywords": metadata.get("global_top_keywords", [])[:50],
        "global_product_name_candidates": metadata.get("global_product_name_candidates", [])[:80],
        "documents": doc_summaries,
    }

    return (
        "You are a technical writing quality auditor. Analyze consistency across a set of documents. "
        "Identify cross-document issues such as inconsistent product names, terminology drift, "
        "style mismatch, duplicated explanations, and gaps in topic coverage.\n\n"
        "Return STRICT JSON with this schema:\n"
        "{\n"
        "  \"executive_summary\": string,\n"
        "  \"findings\": [\n"
        "    {\n"
        "      \"title\": string,\n"
        "      \"severity\": \"high\"|\"medium\"|\"low\",\n"
        "      \"description\": string,\n"
        "      \"evidence\": [string],\n"
        "      \"affected_documents\": [string],\n"
        "      \"recommendations\": [string]\n"
        "    }\n"
        "  ],\n"
        "  \"normalization_map\": [\n"
        "    {\"concept\": string, \"preferred_term\": string, \"variants\": [string]}\n"
        "  ],\n"
        "  \"action_plan\": [string]\n"
        "}\n\n"
        "Use only provided evidence. If uncertain, mark in description.\n\n"
        f"Metadata:\n{json.dumps(compact_metadata, ensure_ascii=True)}"
    )


def _extract_json_object(content: str) -> dict:
    content = content.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", content, re.DOTALL)
    if fenced:
        return json.loads(fenced.group(1))

    start = content.find("{")
    end = content.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("No JSON object found in LLM response.")
    return json.loads(content[start : end + 1])


def run_consistency_analysis(metadata: dict, config: dict) -> dict:
    provider = normalize_provider(config.get("llm_provider", OLLAMA_PROVIDER))
    model_name = resolve_model_name(config)
    if not model_name:
        raise RuntimeError("No model name configured. Set LLM Model in configuration.")

    client = create_client(provider, config)
    prompt = build_llm_prompt(metadata)
    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=min(float(config.get("llm_temperature", 0.1)), 0.3),
        max_tokens=max(int(config.get("llm_max_tokens", 2000)), 2000),
    )

    content = (response.choices[0].message.content or "").strip()
    result = _extract_json_object(content)
    result["_model_used"] = model_name
    result["_provider_used"] = provider
    return result


def write_analysis_docx(analysis: dict, metadata: dict, output_docx_path: Path) -> None:
    doc = Document()
    doc.add_heading("Cross-Document Consistency Analysis", level=1)

    generated = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    doc.add_paragraph(f"Generated at: {generated}")
    doc.add_paragraph(f"Input folder: {metadata.get('input_folder', '')}")
    doc.add_paragraph(f"Document count: {metadata.get('document_count', 0)}")
    doc.add_paragraph(f"Model used: {analysis.get('_model_used', '')} ({analysis.get('_provider_used', '')})")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(str(analysis.get("executive_summary", "No summary provided.")))

    doc.add_heading("Findings", level=2)
    findings = analysis.get("findings", [])
    if not findings:
        doc.add_paragraph("No findings returned by the model.")
    for idx, finding in enumerate(findings, start=1):
        title = finding.get("title", f"Finding {idx}")
        severity = finding.get("severity", "unknown")
        doc.add_heading(f"{idx}. {title} [{severity}]", level=3)
        doc.add_paragraph(str(finding.get("description", "")))

        evidence = finding.get("evidence", [])
        if evidence:
            doc.add_paragraph("Evidence:")
            for item in evidence:
                doc.add_paragraph(str(item), style="List Bullet")

        affected = finding.get("affected_documents", [])
        if affected:
            doc.add_paragraph("Affected documents:")
            for item in affected:
                doc.add_paragraph(str(item), style="List Bullet")

        recs = finding.get("recommendations", [])
        if recs:
            doc.add_paragraph("Recommendations:")
            for item in recs:
                doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("Terminology Normalization Map", level=2)
    normalization_map = analysis.get("normalization_map", [])
    if not normalization_map:
        doc.add_paragraph("No normalization map provided.")
    for row in normalization_map:
        concept = row.get("concept", "")
        preferred = row.get("preferred_term", "")
        variants = ", ".join(row.get("variants", []))
        doc.add_paragraph(f"Concept: {concept}")
        doc.add_paragraph(f"Preferred term: {preferred}")
        doc.add_paragraph(f"Variants found: {variants}")

    doc.add_heading("Recommended Action Plan", level=2)
    action_plan = analysis.get("action_plan", [])
    if not action_plan:
        doc.add_paragraph("No action plan provided.")
    else:
        for step in action_plan:
            doc.add_paragraph(str(step), style="List Number")

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate AI-assisted cross-document consistency analysis DOCX from metadata JSON."
    )
    parser.add_argument(
        "--metadata-json",
        required=True,
        help="Path to metadata JSON produced by consistency_metadata_tool.py",
    )
    parser.add_argument(
        "--output-docx",
        default="output/consistency/consistency_analysis.docx",
        help="Output DOCX path for the analysis report.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    metadata_path = Path(args.metadata_json).resolve()
    output_docx = Path(args.output_docx).resolve()

    if not metadata_path.exists():
        raise SystemExit(f"Metadata JSON not found: {metadata_path}")

    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    config = load_config()
    analysis = run_consistency_analysis(metadata, config)
    write_analysis_docx(analysis, metadata, output_docx)

    print("Consistency analysis complete.")
    print(f"Output report: {output_docx}")


if __name__ == "__main__":
    main()
