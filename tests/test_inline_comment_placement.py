import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

from toolkit.document_processor import _apply_inline_corrections_to_paragraph, _filter_comment_explanation


class InlineCommentPlacementTests(unittest.TestCase):
    def test_inline_comments_are_inserted_at_sentence_end(self):
        doc = Document()
        para = doc.add_paragraph("Group is a container for monitoring assets")
        content = para.text

        corrections = [
            {
                "original": "Group",
                "corrected": "A group",
                "explanation": "Awkward phrasing.",
                "preferred_start": 0,
            }
        ]

        _apply_inline_corrections_to_paragraph(
            para,
            content,
            corrections,
            {
                "highlight_corrections": False,
                "show_deletion_markers": True,
                "add_comments": True,
            },
        )

        text = para.text
        self.assertEqual(
            "A group is a container for monitoring assets [Awkward phrasing.]",
            text,
        )

    def test_inline_comments_deduplicate_repeated_explanations(self):
        doc = Document()
        para = doc.add_paragraph("This is a senttense with duplicate explanation triggers")
        content = para.text

        corrections = [
            {
                "original": "senttense",
                "corrected": "sentence",
                "explanation": "Spelling error in 'senttense'.",
                "preferred_start": content.find("senttense"),
            },
            {
                "original": "triggers",
                "corrected": "triggers",
                "explanation": "Spelling error in 'senttense'.",
                "preferred_start": content.find("triggers"),
            },
        ]

        _apply_inline_corrections_to_paragraph(
            para,
            content,
            corrections,
            {
                "highlight_corrections": False,
                "show_deletion_markers": True,
                "add_comments": True,
            },
        )

        text = para.text
        self.assertEqual(1, text.count("[Spelling error in 'senttense'.]"))

    def test_terminal_punctuation_comments_can_be_suppressed(self):
        doc = Document()
        para = doc.add_paragraph("Install package")
        content = para.text

        corrections = [
            {
                "original": "Install package",
                "corrected": "Install package.",
                "explanation": "Missing terminal period in sentence-style list item.",
                "preferred_start": 0,
            }
        ]

        _apply_inline_corrections_to_paragraph(
            para,
            content,
            corrections,
            {
                "highlight_corrections": False,
                "show_deletion_markers": True,
                "add_comments": True,
                "notify_terminal_punctuation": False,
            },
        )

        self.assertEqual("Install package.", para.text)
        self.assertNotIn("[Missing terminal period in sentence-style list item.]", para.text)

    def test_custom_terminal_punctuation_suppress_strings_are_applied(self):
        explanation = "Model says sentence-style list item needs terminal punctuation"

        with TemporaryDirectory() as tmp_dir:
            suppress_file = Path(tmp_dir) / "terminal_punctuation_suppress_strings.txt"
            suppress_file.write_text(
                "sentence-style list item needs terminal punctuation\n",
                encoding="utf-8",
            )
            with patch(
                "toolkit.document_processor.TERMINAL_PUNCTUATION_SUPPRESS_STRINGS_PATH",
                suppress_file,
            ):
                filtered = _filter_comment_explanation(
                    explanation,
                    {
                        "notify_terminal_punctuation": False,
                    },
                )

        self.assertEqual("", filtered)


if __name__ == "__main__":
    unittest.main()
